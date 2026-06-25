#!/usr/bin/env python3
"""生成 final 文件夹使用说明 Word 文档。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parent / "香薰机实体设备使用手册.docx"


def set_cn_font(run, name="宋体", size=11, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_title(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_cn_font(run, "黑体", 16 if level == 1 else 14, bold=True)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        set_cn_font(run)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        set_cn_font(run)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_cn_font(run, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # 封面
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("智能香薰机（ESP32）\n使用与配置手册")
    set_cn_font(r, "黑体", 22, True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("人车共情闭环系统 · final 文件夹说明\n（面向非技术用户）")
    set_cn_font(r2, "宋体", 12)

    doc.add_paragraph()

    add_title(doc, "一、这份文档是做什么的？", 1)
    add_para(
        doc,
        "本手册说明项目 final 文件夹中的内容：如何用电脑给香薰机烧录程序、"
        "如何连 WiFi 和 MQTT、如何与「人车共情闭环系统」一起工作，"
        "以及如何用网页或手机单独控制香薰机。按步骤操作即可，不要求编程基础。",
    )

    add_title(doc, "二、final 文件夹里有什么？", 1)
    add_table(
        doc,
        ["文件名", "作用", "是否需要修改"],
        [
            ["final.ino", "香薰机主程序（烧录到 ESP32）", "一般不用改，除非改引脚"],
            ["config.h", "WiFi 名称、密码、MQTT 地址", "必须按你的环境填写"],
            ["aroma_control.html", "手机/电脑控制网页", "用浏览器打开即可"],
            ["open_firewall_mqtt.bat", "Windows 防火墙放行 1883 端口", "右键管理员运行一次"],
            ["gen_firmware.py", "开发者重新生成固件用", "普通用户可忽略"],
            ["final.ino.bak", "旧版程序备份", "可忽略"],
        ],
    )

    add_title(doc, "三、系统整体是怎么连起来的？", 1)
    add_para(doc, "可以把它理解成三条线：", bold=True)
    add_bullet(doc, "电脑运行「人车共情系统」（python run.py）→ 通过 MQTT 发命令")
    add_bullet(doc, "香薰机（ESP32）连同一 WiFi → 接收 MQTT 命令 → 喷香、变色、上报状态")
    add_bullet(doc, "你也可以用手机/电脑打开 aroma_control.html → 直接控制香薰机（不经过主系统）")
    add_para(
        doc,
        "MQTT 相当于「对讲频道」：电脑往 cabin/aroma/control 发指令，"
        "香薰机执行后往 cabin/aroma/status 回报是否在喷、进度多少。",
    )

    add_title(doc, "四、你需要准备什么？", 1)
    add_title(doc, "4.1 硬件", 2)
    add_bullet(doc, "ESP32 开发板（已焊接四路喷香 + WS2812 灯带）")
    add_bullet(doc, "USB 数据线（必须是数据线，不能只是充电线）")
    add_bullet(doc, "Windows 电脑（与香薰机、MQTT 在同一局域网 WiFi 下）")

    add_title(doc, "4.2 软件（按顺序安装）", 2)
    add_number(doc, "Arduino IDE 2.x（https://www.arduino.cc/en/software）")
    add_number(doc, "在 Arduino IDE 中安装「ESP32 开发板支持」（开发板管理器搜索 esp32）")
    add_number(doc, "安装以下库（库管理器搜索）：PubSubClient、ArduinoJson、Adafruit NeoPixel")
    add_number(doc, "安装 Mosquitto MQTT（Windows 安装包，作为消息中转站）")
    add_number(doc, "安装 Python 3.10+，并在项目根目录执行：pip install -r requirements.txt")

    add_title(doc, "五、第一次使用：分步操作", 1)

    add_title(doc, "步骤 1：填写 config.h", 2)
    add_para(doc, "用记事本或 VS Code 打开 final/config.h，修改以下几行：")
    add_table(
        doc,
        ["配置项", "填什么", "示例"],
        [
            ["WIFI_SSID", "你的 WiFi 名称", "test"],
            ["WIFI_PASSWORD", "WiFi 密码", "（你的密码）"],
            ["MQTT_BROKER", "运行 Mosquitto 的电脑的局域网 IP", "192.168.117.94"],
            ["MQTT_PORT", "一般不用改", "1883"],
            ["TOPIC_PREFIX / DEVICE_NAME", "一般不用改", "cabin / aroma"],
        ],
    )
    add_para(
        doc,
        "重要：MQTT_BROKER 不能填 localhost，必须填电脑在 WiFi 下的 IP。"
        "在电脑上打开 cmd，输入 ipconfig，找「无线局域网 IPv4 地址」。",
    )

    add_title(doc, "步骤 2：烧录程序到 ESP32", 2)
    add_number(doc, "USB 连接 ESP32 到电脑（建议暂时断开灯带和大功率负载，避免烧录失败）")
    add_number(doc, "Arduino IDE 打开 final/final.ino")
    add_number(doc, "工具 → 开发板 → ESP32 Dev Module")
    add_number(doc, "工具 → 上传速度 → 115200")
    add_number(doc, "工具 → 端口 → 选择 COM 口（如 COM7）")
    add_number(doc, "点击「上传」，等待完成")
    add_para(doc, "串口监视器设为 115200，上电后应看到「成功连接 WiFi」和「成功连接 MQTT Broker」。")

    add_title(doc, "步骤 3：放行 Windows 防火墙（必做）", 2)
    add_para(
        doc,
        "ESP32 从 WiFi 访问电脑的 MQTT 时，Windows 默认会拦截。"
        "请右键 final/open_firewall_mqtt.bat → 以管理员身份运行，看到成功提示即可。",
    )

    add_title(doc, "步骤 4：启动 Mosquitto 和主系统", 2)
    add_number(doc, "确保 Mosquitto 服务正在运行（端口 1883）")
    add_number(doc, "打开 PowerShell，进入项目根目录")
    add_number(doc, "执行：$env:SKIP_DEVICE_SIMS=\"aroma\"  然后  python run.py")
    add_para(
        doc,
        "SKIP_DEVICE_SIMS=aroma 表示：不用电脑模拟香薰，改用实体 ESP32。"
        "否则会两个程序抢同一个 MQTT 频道。",
    )
    add_number(doc, "浏览器打开 http://127.0.0.1:5000 查看设备状态，香薰应显示「在线」")

    add_title(doc, "步骤 5：用网页单独控制香薰机（可选）", 2)
    add_number(doc, "双击打开 final/aroma_control.html")
    add_number(doc, "填入 ESP32 串口日志里的 IP 地址，点「保存」")
    add_number(doc, "选择模式（平静/舒缓/提油等），调节速度，点「执行选中模式」")
    add_para(doc, "手机也可使用：手机与香薰机连同一 WiFi，用浏览器打开该 html 文件即可。")

    add_title(doc, "六、香薰机有哪些模式？", 1)
    add_para(doc, "与主系统对接的三种（AI 会自动下发）：", bold=True)
    add_table(
        doc,
        ["模式", "通道/气味", "总时长", "喷洒方式"],
        [
            ["平静", "NH 洋甘菊薰衣草", "90 秒", "先爆发喷 6 秒，再间歇 3.5s开/6.5s停"],
            ["舒缓", "PL 白玉兰", "75 秒", "先爆发 7 秒，再间歇 4s开/5s停"],
            ["提神", "NL 薄荷", "50 秒", "先爆发 8 秒，再持续长喷 9s开/1.5s停"],
            ["关闭", "—", "—", "停止喷香，橙色呼吸待机灯"],
        ],
    )
    add_para(doc, "网页上还可用的扩展模式：专注、共情、振奋、冥想、净化、浪漫、深度放松、派对、夜灯等。")

    add_title(doc, "七、硬件通道说明", 1)
    add_table(
        doc,
        ["代码", "引脚", "对应香型", "灯效"],
        [
            ["PH", "GPIO 25", "青柠罗勒柑橘", "青绿色"],
            ["PL", "GPIO 26", "白玉兰", "暖白色"],
            ["NH", "GPIO 27", "洋甘菊薰衣草", "紫色"],
            ["NL", "GPIO 14", "薄荷", "冰蓝色"],
        ],
    )
    add_para(doc, "灯带接 GPIO 13，16 颗 WS2812。")

    add_title(doc, "八、常见问题", 1)
    add_table(
        doc,
        ["现象", "可能原因", "解决办法"],
        [
            ["MQTT 连接失败 rc=-2", "防火墙或 IP 不对", "运行 open_firewall_mqtt.bat；核对 config.h 中 IP"],
            ["网页香薰一直「等待设备」", "ESP32 未连 MQTT 或模拟器冲突", "看串口是否连上；启动时 SKIP_DEVICE_SIMS=aroma"],
            ["COM 口断断续续", "USB 线差或电流过大", "换数据线/后置 USB；喷香时独立 5V 供电"],
            ["烧录失败", "上传速度太高", "上传速度改 115200；按住 BOOT 再上传"],
            ["平静/舒缓没味道", "未重新烧录新固件", "烧录最新 final.ino，模式含爆发喷"],
        ],
    )

    add_title(doc, "九、日常启动清单（可打印）", 1)
    add_number(doc, "Mosquitto 已启动")
    add_number(doc, "ESP32 上电，串口见 MQTT 已连接")
    add_number(doc, "PowerShell：$env:SKIP_DEVICE_SIMS=\"aroma\"")
    add_number(doc, "python run.py")
    add_number(doc, "浏览器 http://127.0.0.1:5000")

    add_title(doc, "十、获取帮助时应提供的信息", 1)
    add_bullet(doc, "config.h 中的 WiFi 名和 MQTT IP（不要发密码）")
    add_bullet(doc, "串口监视器 115200 的 WiFi/MQTT 相关日志")
    add_bullet(doc, "电脑 ipconfig 的 IPv4 地址")
    add_bullet(doc, "网页 /device_status 或 aroma_control 显示的状态")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— 文档随 final 文件夹程序更新 —")
    set_cn_font(r, size=10)

    doc.save(OUT)
    print(f"已生成: {OUT}")


if __name__ == "__main__":
    build()
